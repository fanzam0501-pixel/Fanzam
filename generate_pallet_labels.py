#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量生成栈板标签文档
基于模板，填充 SN 码，每个 SN 生成一个标签页
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def extract_sn_from_filename(filename):
    """从文件名提取 SN 码"""
    match = re.search(r'SN:5504AJML\d+', filename)
    if match:
        return match.group(0)
    return None

def generate_pallet_labels(template_path, barcodes_dir, output_path, carton_qty=25):
    """
    批量生成栈板标签
    
    Args:
        template_path: 模板文档路径
        barcodes_dir: 条形码图片目录
        output_path: 输出文档路径
        carton_qty: 每箱数量（默认25）
    """
    # 获取所有条形码文件并排序
    barcode_files = []
    for f in os.listdir(barcodes_dir):
        if f.endswith('.png') and f.startswith('SN:'):
            barcode_files.append(f)
    
    # 按数字排序
    barcode_files.sort(key=lambda x: int(re.search(r'\d+', x).group()))
    
    print(f"找到 {len(barcode_files)} 个条形码")
    
    # 创建新文档
    doc = Document()
    
    # 设置页面边距（根据实际情况调整）
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
    
    total_sn = len(barcode_files)
    
    for idx, barcode_file in enumerate(barcode_files, 1):
        sn_code = extract_sn_from_filename(barcode_file)
        if not sn_code:
            continue
        
        # 添加分页（除了第一页）
        if idx > 1:
            doc.add_page_break()
        
        # 添加标签内容
        # COLOUR
        p = doc.add_paragraph()
        p.add_run(f"COLOUR：WHITE").bold = True
        
        # MODEL
        p = doc.add_paragraph()
        p.add_run(f"MODEL：HF-01").bold = True
        
        # SKU/ITEM（留空）
        p = doc.add_paragraph()
        p.add_run("SKU/ITEM：").bold = True
        
        # QTY
        p = doc.add_paragraph()
        p.add_run(f"QTY：{carton_qty}").bold = True
        
        # G.W.
        p = doc.add_paragraph()
        p.add_run("G.W.:  KG").bold = True
        
        # Carton NO.
        p = doc.add_paragraph()
        p.add_run(f"Carton NO. {idx} / {total_sn}").bold = True
        
        # S/N(TOTAL)
        p = doc.add_paragraph()
        p.add_run(f"S/N(TOTAL): {sn_code}").bold = True
        
        # 插入条形码图片
        barcode_path = os.path.join(barcodes_dir, barcode_file)
        if os.path.exists(barcode_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(barcode_path, width=Cm(6))
        
        # MADE IN CHINA
        p = doc.add_paragraph()
        p.add_run("MADE IN CHINA").bold = True
        
        # 分隔线
        if idx < total_sn:
            p = doc.add_paragraph()
            p.add_run("_" * 50)
        
        print(f"[{idx}/{total_sn}] 已添加标签: {sn_code}")
    
    # 保存文档
    doc.save(output_path)
    print(f"\n✅ 文档已保存: {output_path}")
    print(f"共生成 {total_sn} 个栈板标签")
    return output_path

def generate_pallet_labels_from_template(template_path, barcodes_dir, output_path):
    """
    基于模板文档结构批量生成标签（保留模板格式）
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy
    
    # 获取所有条形码文件并排序
    barcode_files = []
    for f in os.listdir(barcodes_dir):
        if f.endswith('.png') and f.startswith('SN:'):
            barcode_files.append(f)
    
    barcode_files.sort(key=lambda x: int(re.search(r'\d+', x).group()))
    print(f"找到 {len(barcode_files)} 个条形码")
    
    # 打开模板
    template_doc = Document(template_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 复制模板样式
    for style in template_doc.styles:
        try:
            new_doc.styles.add_style(style.name, style.type)
        except:
            pass
    
    total_sn = len(barcode_files)
    
    for idx, barcode_file in enumerate(barcode_files, 1):
        sn_code = extract_sn_from_filename(barcode_file)
        if not sn_code:
            continue
        
        # 添加分页（除了第一页）
        if idx > 1:
            new_doc.add_page_break()
        
        # 复制模板内容并替换 SN 码
        for para in template_doc.paragraphs:
            new_para = new_doc.add_paragraph()
            
            # 复制段落样式
            new_para.style = para.style
            new_para.alignment = para.alignment
            
            # 复制并修改文本
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
            
            # 替换 S/N 相关内容
            if 'S/N' in new_para.text or 'SN' in new_para.text:
                # 替换 S/N(TOTAL): 后的内容
                if 'S/N(TOTAL):' in new_para.text:
                    new_para.clear()
                    new_run = new_para.add_run(f"S/N(TOTAL): {sn_code}")
                    new_run.bold = True
            
            # 替换 Carton NO.
            if 'Carton NO.' in new_para.text:
                new_para.clear()
                new_run = new_para.add_run(f"Carton NO. {idx} / {total_sn}")
                new_run.bold = True
            
            # 替换 QTY
            if 'QTY：' in new_para.text or 'QTY:' in new_para.text:
                new_para.clear()
                new_run = new_para.add_run("QTY：25")
                new_run.bold = True
        
        # 在文档末尾添加条形码图片
        barcode_path = os.path.join(barcodes_dir, barcode_file)
        if os.path.exists(barcode_path):
            p = new_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(barcode_path, width=Cm(6))
        
        print(f"[{idx}/{total_sn}] 已添加标签: {sn_code}")
    
    # 保存文档
    new_doc.save(output_path)
    print(f"\n✅ 文档已保存: {output_path}")
    print(f"共生成 {total_sn} 个栈板标签")
    return output_path

# ==================== 主程序 ====================
if __name__ == "__main__":
    template_path = "/root/.openclaw/media/inbound/栈板SN码-英规---425a57a6-b194-4415-986a-d5e405ed80ee.docx"
    barcodes_dir = "/root/.openclaw/media/inbound/barcodes_new"
    output_path = "/root/.openclaw/workspace/栈板标签_全部SN码.docx"
    
    # 使用基于模板的方法
    print("开始生成栈板标签...")
    print(f"模板: {template_path}")
    print(f"条形码目录: {barcodes_dir}")
    print()
    
    generate_pallet_labels_from_template(template_path, barcodes_dir, output_path)
