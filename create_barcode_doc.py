#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建 Word 文档，包含条形码图片表格
表格：每页3行2列，单元格尺寸 6×7cm
每个单元格包含：资源图片（上）+ 条形码图片（下）
"""

from docx import Document
from docx.shared import Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_size(cell, width_cm, height_cm):
    """设置单元格尺寸"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 567 twips = 1cm
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    # 设置行高
    tr = tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def add_centered_image(cell, image_path, width_cm=None):
    """在单元格中添加居中图片"""
    if not os.path.exists(image_path):
        print(f"图片不存在: {image_path}")
        return False
    
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run()
    if width_cm:
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        run.add_picture(image_path)
    return True

def create_barcode_document():
    # 路径设置
    base_path = "/root/.openclaw/media/inbound/extracted_1"
    barcodes_path = os.path.join(base_path, "barcodes")
    resource_image = os.path.join(base_path, "资源 4@4x.png")
    
    # 获取所有条形码图片并按顺序排序
    barcode_files = []
    for f in os.listdir(barcodes_path):
        if f.endswith('.png') and f.startswith('SN:'):
            barcode_files.append(f)
    
    # 按文件名中的数字排序
    barcode_files.sort(key=lambda x: int(x.split('SN:5504AJML2644')[1].split('.')[0]))
    
    print(f"找到 {len(barcode_files)} 张条形码图片")
    
    # 创建文档
    doc = Document()
    
    # 设置页面边距，确保表格能放下
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
    
    # 每页6张图片（3行2列）
    images_per_page = 6
    total_images = len(barcode_files)
    
    for page_num in range(0, total_images, images_per_page):
        # 添加分页（除了第一页）
        if page_num > 0:
            doc.add_page_break()
        
        # 创建表格：3行2列
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # 处理这一页的6张图片
        for row_idx in range(3):
            for col_idx in range(2):
                img_idx = page_num + row_idx * 2 + col_idx
                if img_idx >= total_images:
                    break
                
                cell = table.cell(row_idx, col_idx)
                barcode_file = barcode_files[img_idx]
                barcode_path = os.path.join(barcodes_path, barcode_file)
                
                # 设置单元格尺寸 6×7cm
                set_cell_size(cell, 6, 7)
                
                # 先添加资源图片（上方）
                add_centered_image(cell, resource_image, width_cm=5)
                
                # 在同一单元格内添加条形码图片（下方）
                # 先添加一个换行
                cell.paragraphs[0].add_run().add_break()
                # 添加条形码图片，宽度5cm
                add_centered_image(cell, barcode_path, width_cm=5)
                
                print(f"已添加: {barcode_file}")
    
    # 保存文档
    output_path = "/root/.openclaw/workspace/条形码文档.docx"
    doc.save(output_path)
    print(f"\n文档已保存: {output_path}")
    print(f"共 {total_images} 张条形码图片，每页6张，共 {(total_images // 6) + (1 if total_images % 6 else 0)} 页")
    return output_path

if __name__ == "__main__":
    create_barcode_document()
