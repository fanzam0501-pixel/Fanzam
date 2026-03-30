#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
栈板标签文档生成 - 格式：
- 每页上方：固定标签信息（COLOUR/MODEL/QTY等）
- 横线下方：25个二维码（5行×5列）
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
from PIL import Image

def set_cell_size(cell, width_cm, height_cm):
    """设置单元格尺寸"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    tr = tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def copy_paragraph(source_para, target_doc):
    """复制段落到新文档"""
    new_para = target_doc.add_paragraph()
    new_para.style = source_para.style
    new_para.alignment = source_para.alignment
    
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
    
    return new_para

def generate_pallet_labels_with_qrcodes(template_path, barcodes_dir, output_path, qty_per_page=25):
    """
    生成栈板标签文档
    每页：横线上方固定内容 + 横线下方25个二维码
    """
    # 获取所有条形码文件
    barcode_files = []
    for f in os.listdir(barcodes_dir):
        if f.endswith('.png') and f.startswith('SN:'):
            barcode_files.append(f)
    
    barcode_files.sort(key=lambda x: int(re.search(r'\d+', x).group()))
    total = len(barcode_files)
    
    print(f"找到 {total} 个条形码")
    
    # 打开模板
    template_doc = Document(template_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置页面边距
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
    
    # 计算需要多少页
    pages_needed = (total + qty_per_page - 1) // qty_per_page
    
    page_num = 0
    idx = 0
    
    while idx < total:
        page_num += 1
        
        # 添加分页（除了第一页）
        if page_num > 1:
            new_doc.add_page_break()
        
        # 复制横线上方的内容（段落 0-17）
        for para_idx in range(18):  # 0-17 包含横线
            source_para = template_doc.paragraphs[para_idx]
            new_para = copy_paragraph(source_para, new_doc)
        
        # 在横线下方添加空行
        new_doc.add_paragraph()
        
        # 创建二维码表格（5行×5列 = 25个）
        rows = 5
        cols = 5
        table = new_doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充二维码
        for row_idx in range(rows):
            for col_idx in range(cols):
                cell = table.cell(row_idx, col_idx)
                
                # 设置单元格尺寸（3cm × 3cm）
                set_cell_size(cell, 3, 3)
                
                if idx < total:
                    barcode_file = barcode_files[idx]
                    barcode_path = os.path.join(barcodes_dir, barcode_file)
                    
                    # 提取 SN 码作为文字
                    sn_match = re.search(r'SN:5504AJML\d+', barcode_file)
                    sn_code = sn_match.group(0) if sn_match else ""
                    
                    # 添加 SN 文字
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(sn_code[-5:])  # 只显示最后5位
                    run.font.size = Pt(8)
                    
                    # 添加二维码图片
                    if os.path.exists(barcode_path):
                        p.add_run().add_break()
                        run = p.add_run()
                        run.add_picture(barcode_path, width=Cm(2.2))
                    
                    idx += 1
                    print(f"[{idx}/{total}] 已添加: {sn_code}")
                else:
                    # 空白单元格
                    pass
        
        # 添加说明文字
        new_doc.add_paragraph()
        p = new_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"第 {page_num} 页 / 共 {pages_needed} 页 | 本页: {(page_num-1)*qty_per_page+1}-{min(page_num*qty_per_page, total)}")
        run.font.size = Pt(9)
        run.italic = True
    
    # 保存文档
    new_doc.save(output_path)
    print(f"\n✅ 文档已保存: {output_path}")
    print(f"共 {total} 个二维码，{pages_needed} 页，每页最多 {qty_per_page} 个")
    return output_path

if __name__ == "__main__":
    template_path = "/root/.openclaw/media/inbound/栈板SN码-英规---425a57a6-b194-4415-986a-d5e405ed80ee.docx"
    barcodes_dir = "/root/.openclaw/media/inbound/barcodes_new"
    output_path = "/root/.openclaw/workspace/栈板标签_二维码版.docx"
    
    print("开始生成栈板标签文档...")
    print(f"模板: {template_path}")
    print(f"条形码目录: {barcodes_dir}")
    print(f"每页二维码数量: 25个 (5×5)")
    print()
    
    generate_pallet_labels_with_qrcodes(template_path, barcodes_dir, output_path, qty_per_page=25)
