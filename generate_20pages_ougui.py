#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
复制欧规模板第一页，创建20页
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

def copy_paragraph(source_para, target_doc):
    """复制段落到新文档"""
    new_para = target_doc.add_paragraph()
    new_para.style = source_para.style
    new_para.alignment = source_para.alignment
    
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        if run.font.color and run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
    
    return new_para

def copy_table(source_table, target_doc):
    """复制表格到新文档"""
    rows = len(source_table.rows)
    cols = len(source_table.rows[0].cells) if rows > 0 else 0
    
    new_table = target_doc.add_table(rows=rows, cols=cols)
    new_table.style = source_table.style
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_cell = new_table.cell(i, j)
            new_cell.text = cell.text
            # 复制单元格样式（简化处理）
            for para in cell.paragraphs:
                if para.text.strip():
                    new_para = new_cell.paragraphs[0]
                    new_para.text = para.text
                    new_para.alignment = para.alignment
    
    return new_table

def generate_20_pages(template_path, output_path):
    """
    复制模板第一页内容，创建20页
    """
    # 打开模板
    template_doc = Document(template_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置页面边距
    sections = new_doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
    
    # 复制20次
    for page_num in range(1, 21):
        # 添加分页（除了第一页）
        if page_num > 1:
            new_doc.add_page_break()
        
        # 复制所有段落
        for para in template_doc.paragraphs:
            copy_paragraph(para, new_doc)
        
        # 复制所有表格
        for table in template_doc.tables:
            copy_table(table, new_doc)
        
        print(f"第 {page_num}/20 页已创建")
    
    # 保存文档
    new_doc.save(output_path)
    print(f"\n✅ 文档已保存: {output_path}")
    print(f"共 20 页")
    return output_path

if __name__ == "__main__":
    template_path = "/root/.openclaw/media/inbound/栈板SN码-欧规-20页---af1726a6-24b3-48af-9796-259d09009d85.docx"
    output_path = "/root/.openclaw/workspace/栈板SN码-欧规-20页-副本.docx"
    
    print("开始生成20页欧规标签...")
    generate_20_pages(template_path, output_path)
