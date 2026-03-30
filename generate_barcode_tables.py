#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建条形码排列文档
表格分3列，25个一组，按顺序排列
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

def set_cell_size(cell, width_cm, height_cm):
    """设置单元格尺寸"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    tr = tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def generate_barcodes_table(barcodes_dir, output_path, cols=3, per_table=25):
    """
    生成条形码表格文档
    
    Args:
        barcodes_dir: 条形码图片目录
        output_path: 输出文档路径
        cols: 表格列数（默认3列）
        per_table: 每个表格包含的条形码数量（默认25个）
    """
    # 获取所有条形码文件
    barcode_files = []
    for f in os.listdir(barcodes_dir):
        if f.endswith('.png') and f.startswith('SN:'):
            barcode_files.append(f)
    
    barcode_files.sort(key=lambda x: int(re.search(r'\d+', x).group()))
    total = len(barcode_files)
    
    print(f"找到 {total} 个条形码")
    print(f"表格设置: {cols}列, 每表格{per_table}个")
    
    # 创建新文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
    
    # 计算需要多少个表格
    table_count = (total + per_table - 1) // per_table
    
    idx = 0
    table_num = 0
    
    while idx < total:
        table_num += 1
        remaining = total - idx
        current_table_count = min(per_table, remaining)
        
        # 计算当前表格需要多少行
        rows = (current_table_count + cols - 1) // cols
        
        # 添加表格标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"【第 {table_num} 组】")
        run.bold = True
        run.font.size = Pt(12)
        
        # 创建表格
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # 填充表格（限制每组数量）
        count_in_this_table = 0
        for row_idx in range(rows):
            for col_idx in range(cols):
                cell = table.cell(row_idx, col_idx)
                
                # 设置单元格尺寸 (5.2cm × 2cm)
                set_cell_size(cell, 5.2, 2)

                # 只填充指定的数量（每组25个）
                if count_in_this_table < current_table_count and idx < total:
                    barcode_file = barcode_files[idx]
                    barcode_path = os.path.join(barcodes_dir, barcode_file)

                    # 提取 SN 码
                    sn_match = re.search(r'SN:5504AJML(\d+)', barcode_file)
                    sn_num = sn_match.group(1) if sn_match else ""

                    # 添加图片
                    if os.path.exists(barcode_path):
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(barcode_path, width=Cm(5))
                    
                    idx += 1
                    count_in_this_table += 1
                    print(f"[{idx}/{total}] 已添加: SN:5504AJML{sn_num} (第{table_num}组第{count_in_this_table}个)")
                else:
                    # 空白单元格（合并单元格或留空）
                    pass
        
        # 表格之间添加空两行，然后分页
        if idx < total:
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_page_break()
    
    # 添加统计信息
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n总计: {total} 个条形码")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"分 {table_num} 组，每组最多 {per_table} 个，3列排列")
    run.font.size = Pt(11)
    
    # 保存文档
    doc.save(output_path)
    print(f"\n✅ 文档已保存: {output_path}")
    print(f"共 {total} 个条形码，{table_num} 个表格")
    return output_path

if __name__ == "__main__":
    barcodes_dir = "/root/.openclaw/media/inbound/extracted_1/barcodes"
    output_path = "/root/.openclaw/workspace/条形码_3列25个一组.docx"
    
    print("开始生成条形码表格文档...")
    print(f"条形码目录: {barcodes_dir}")
    print()
    
    generate_barcodes_table(barcodes_dir, output_path, cols=3, per_table=25)
